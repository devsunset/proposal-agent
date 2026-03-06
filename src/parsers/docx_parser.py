"""
DOCX 문서 파서

python-docx를 사용해 .docx/.doc 파일에서 텍스트, 테이블, 헤딩 기반 섹션, 메타데이터를 추출합니다.
RFP가 워드 문서로 제공될 때 사용됩니다.
- 단락·테이블을 문서 순서대로 raw_text에 포함 (요구기술·평가기준 등 누락 방지)
- 머리글/바닥글, 문서 구조(섹션 목차), 메타데이터까지 최대한 추출하여 RFP 분석에 활용
"""

from pathlib import Path
from typing import Any, Dict, List

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from .base_parser import BaseParser
from ..utils.logger import get_logger

logger = get_logger("docx_parser")


def _iter_block_items(parent):
    """문서 내 단락(Paragraph)과 테이블(Table)을 문서 순서대로 yield."""
    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_elm = parent.element.body
    else:
        return
    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _table_to_text(table: Table) -> str:
    """테이블을 읽기 쉬운 텍스트(헤더 + 행들)로 변환."""
    lines = []
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    return "\n".join(lines) if lines else ""


class DOCXParser(BaseParser):
    """
    DOCX 문서 전용 파서.

    지원 확장자: .docx, .doc
    """

    @property
    def supported_extensions(self) -> List[str]:
        return [".docx", ".doc"]

    def parse(self, file_path: Path) -> Dict[str, Any]:
        """
        DOCX를 파싱하여 raw_text, tables, sections, metadata, styles, document_structure 반환.
        본문·테이블·머리글/바닥글을 모두 raw_text에 반영하고, 문서 구조(목차)를 별도로 제공.
        """
        logger.info(f"DOCX 파싱 시작: {file_path}")

        doc = Document(file_path)
        body_text = self.extract_text(file_path)
        tables = self.extract_tables(file_path)
        sections = self._extract_sections(doc)
        metadata = self._extract_metadata(doc)
        header_footer = self._extract_headers_footers(doc)
        document_structure = self._build_document_structure(sections)

        # 본문 + 머리글/바닥글을 합쳐 RFP 분석 시 모든 정보 참조 가능하게
        raw_text_parts = []
        if header_footer.strip():
            raw_text_parts.append("[머리글·바닥글]\n" + header_footer.strip())
        raw_text_parts.append(body_text)
        raw_text = "\n\n".join(raw_text_parts)

        result = {
            "raw_text": raw_text,
            "tables": tables,
            "sections": sections,
            "metadata": metadata,
            "styles": self._extract_styles(doc),
            "document_structure": document_structure,
        }

        logger.info(
            f"DOCX 파싱 완료: {len(result['raw_text'])} 문자, "
            f"{len(result['tables'])} 테이블, {len(result['sections'])} 섹션, "
            f"구조 {len(document_structure)}항목"
        )

        return result

    def extract_text(self, file_path: Path) -> str:
        """전체 텍스트 추출 (단락 + 테이블을 문서 순서대로 포함, 요구기술·평가기준 등 누락 방지)"""
        try:
            doc = Document(file_path)
            parts = []
            for block in _iter_block_items(doc):
                if isinstance(block, Paragraph):
                    text = block.text.strip()
                    if text:
                        parts.append(text)
                elif isinstance(block, Table):
                    table_text = _table_to_text(block)
                    if table_text:
                        parts.append("[테이블]\n" + table_text)
            return "\n\n".join(parts)
        except Exception as e:
            logger.error(f"텍스트 추출 실패: {e}")
            return ""

    def extract_tables(self, file_path: Path) -> List[Dict[str, Any]]:
        """테이블 데이터 추출"""
        tables = []

        try:
            doc = Document(file_path)

            for i, table in enumerate(doc.tables):
                table_data = self._table_to_dict(table, i)
                if table_data:
                    tables.append(table_data)
        except Exception as e:
            logger.error(f"테이블 추출 실패: {e}")

        return tables

    def _table_to_dict(self, table: Table, index: int) -> Dict[str, Any]:
        """docx Table 객체를 headers/rows/raw_data 구조의 딕셔너리로 변환. 첫 행을 헤더로 처리."""
        rows = []

        for row in table.rows:
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            rows.append(row_data)

        if not rows:
            return {}

        # 첫 번째 행을 헤더로 처리
        headers = rows[0] if rows else []
        data_rows = rows[1:] if len(rows) > 1 else []

        return {
            "table_index": index,
            "headers": headers,
            "rows": data_rows,
            "raw_data": rows,
        }

    def _extract_headers_footers(self, doc: Document) -> str:
        """머리글·바닥글 텍스트 추출 (RFP 제목·문서번호 등이 있을 수 있음)."""
        parts = []
        try:
            for section in doc.sections:
                for part_name, part in [("머리글", section.header), ("바닥글", section.footer)]:
                    if part is None:
                        continue
                    paras = [p.text.strip() for p in part.paragraphs if p.text.strip()]
                    if paras:
                        parts.append(f"{part_name}: " + " ".join(paras))
        except Exception as e:
            logger.debug("머리글/바닥글 추출 생략: %s", e)
        return "\n".join(parts)

    def _build_document_structure(self, sections: List[Dict[str, Any]]) -> str:
        """섹션 목록으로 문서 구조(목차) 문자열 생성. RFP 분석 시 전체 구성을 참조용으로 전달."""
        if not sections:
            return ""
        lines = []
        for s in sections:
            title = (s.get("title") or "").strip()
            if not title:
                continue
            level = s.get("level", 1)
            indent = "  " * (level - 1)
            lines.append(f"{indent}- {title}")
        return "\n".join(lines)

    def _extract_sections(self, doc: Document) -> List[Dict[str, Any]]:
        """문서 내 Heading 스타일을 기준으로 섹션(title, content, level) 목록 추출."""
        sections = []
        current_section = {"title": "", "content": [], "level": 0}

        for para in doc.paragraphs:
            # 헤딩 스타일 체크
            if para.style and para.style.name.startswith("Heading"):
                # 이전 섹션 저장
                if current_section["content"] or current_section["title"]:
                    sections.append(current_section)

                # 헤딩 레벨 추출
                level = 1
                if para.style.name[-1].isdigit():
                    level = int(para.style.name[-1])

                current_section = {
                    "title": para.text.strip(),
                    "content": [],
                    "level": level,
                    "style": para.style.name,
                }
            else:
                text = para.text.strip()
                if text:
                    current_section["content"].append(text)

        # 마지막 섹션 저장
        if current_section["content"] or current_section["title"]:
            sections.append(current_section)

        return sections

    def _extract_metadata(self, doc: Document) -> Dict[str, Any]:
        """core_properties에서 제목, 작성자, 주제, 키워드, 생성/수정일 추출."""
        try:
            core_props = doc.core_properties
            return {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "subject": core_props.subject or "",
                "keywords": core_props.keywords or "",
                "created": str(core_props.created) if core_props.created else "",
                "modified": str(core_props.modified) if core_props.modified else "",
            }
        except Exception as e:
            logger.warning(f"메타데이터 추출 실패: {e}")
            return {}

    def _extract_styles(self, doc: Document) -> Dict[str, Any]:
        """문서에서 사용된 단락 스타일 이름 목록 추출 (styles_used)."""
        styles_used = set()

        for para in doc.paragraphs:
            if para.style:
                styles_used.add(para.style.name)

        return {"styles_used": list(styles_used)}
